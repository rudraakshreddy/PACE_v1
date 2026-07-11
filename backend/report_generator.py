"""
Report Generator for PACE Calculation Module
Produces a multi-page WAVE-style .docx report with tables AND embedded charts.
Reference layout: calc_module_ref_pics/RO_report_summary.pdf
"""

import os
import io
import tempfile
from datetime import datetime

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# ─── Colour constants ────────────────────────────────────────────────────────
BRAND_COLOR = RGBColor(65, 105, 225)  # Permionics Blue
RED_WARN    = RGBColor(192, 0, 0)
BLACK       = RGBColor(0, 0, 0)
HEX_BORDER  = "7F7F7F"
CHART_BRAND = "#4169E1"
CHART_BLUE  = "#1565C0"
CHART_ORANGE= "#E65100"
CHART_GRAY  = "#607D8B"

# ─── Matplotlib style ────────────────────────────────────────────────────────
plt.rcParams.update({
    'font.family': 'DejaVu Sans', 'font.size': 8,
    'axes.titlesize': 9, 'axes.titleweight': 'bold',
    'axes.labelsize': 8, 'axes.spines.top': False, 'axes.spines.right': False,
    'axes.grid': True, 'grid.alpha': 0.3, 'figure.dpi': 150,
    'figure.facecolor': 'white',
})


# ─────────────────────────────────────────────────────────────────────────────
# XML / docx helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)

def conv_flow(val, units_dict):
    to_unit = units_dict.get('flow', 'm3/h')
    if to_unit == 'm3/d': return val * 24.0
    if to_unit == 'gpm': return val * 4.40287
    if to_unit == 'gpd': return val * 6340.13
    return val

def conv_press(val, units_dict):
    to_unit = units_dict.get('pressure', 'bar')
    if to_unit == 'psi': return val * 14.5038
    if to_unit == 'MPa': return val * 0.1
    if to_unit == 'kPa': return val * 100.0
    return val

def conv_flux(val, units_dict):
    to_unit = units_dict.get('flux', 'LMH')
    if to_unit == 'GFD': return val * 0.589
    return val

def lbl_flow(units_dict):
    u = units_dict.get('flow', 'm3/h')
    if u == 'm3/h': return 'm³/h'
    if u == 'm3/d': return 'm³/d'
    return u

def lbl_press(units_dict):
    return units_dict.get('pressure', 'bar')

def lbl_flux(units_dict):
    return units_dict.get('flux', 'LMH')

def _set_cell_border(cell, **kw):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    bdr  = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        data = kw.get(edge)
        if data:
            el = OxmlElement(f'w:{edge}')
            for k, v in data.items():
                el.set(qn(f'w:{k}'), str(v))
            bdr.append(el)
    tcPr.append(bdr)

def _thin(color=HEX_BORDER):
    s = {"sz": "4", "val": "single", "color": color}
    return {"top": s, "left": s, "bottom": s, "right": s}

def _no_border():
    s = {"sz": "0", "val": "none", "color": "auto"}
    return {"top": s, "left": s, "bottom": s, "right": s}

def _set_cell_bg(cell, hex_col):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_col}" w:color="auto" w:val="clear"/>')
    tcPr.append(shd)

def _cell_text(cell, text, bold=False, italic=False, size=9, color=BLACK,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    r.font.name = 'Calibri'; r.font.size = Pt(size)

def _run(para, text, bold=False, size=9.5, color=BLACK, italic=False):
    r = para.add_run(str(text))
    r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    r.font.name = 'Calibri'; r.font.size = Pt(size)
    return r

def _col_w(table, ci, cm_val):
    for row in table.rows:
        row.cells[ci].width = Cm(cm_val)

def _spacer(doc, pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)

def _heading(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    _run(p, text, bold=True, size=size, color=BRAND_COLOR)

def _hr(doc):
    """Thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '7F7F7F')
    pBdr.append(bot); pPr.append(pBdr)


# ─────────────────────────────────────────────────────────────────────────────
# Generic table builders
# ─────────────────────────────────────────────────────────────────────────────

def _prop_table(doc, rows, widths=(7.5, 6.5)):
    """2-column property table (Label | Value), thin borders, shaded labels."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (lbl, val) in enumerate(rows):
        # First column: Label
        c0 = tbl.rows[ri].cells[0]
        _cell_text(c0, lbl, size=9)
        _set_cell_border(c0, **_thin("D1D5DB")) # subtle border
        _set_cell_bg(c0, "F8FAFC") # very light slate for label

        # Second column: Value
        c1 = tbl.rows[ri].cells[1]
        _cell_text(c1, val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_border(c1, **_thin("D1D5DB"))
    _col_w(tbl, 0, widths[0]); _col_w(tbl, 1, widths[1])
    return tbl


def _data_table(doc, headers, rows, col_widths=None):
    """
    Standard data table: bold centred header row(s) with shading, then alternating data rows.
    """
    n_cols = max(len(r) for r in headers + rows) if (headers + rows) else 1
    total  = len(headers) + len(rows)
    tbl    = doc.add_table(rows=total, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ri, hrow in enumerate(headers):
        for ci, val in enumerate(hrow[:n_cols]):
            c = tbl.rows[ri].cells[ci]
            _cell_text(c, val or '', bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(30,41,59))
            _set_cell_border(c, **_thin("94A3B8")) # darker border for header
            _set_cell_bg(c, "F1F5F9") # slate-100 header background

    for ri, drow in enumerate(rows):
        r_obj = tbl.rows[len(headers) + ri]
        bg_color = "F8FAFC" if ri % 2 == 1 else "FFFFFF" # alternating rows
        for ci, val in enumerate(drow[:n_cols]):
            c = r_obj.cells[ci]
            al = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _cell_text(c, val if val is not None else '', size=8.5, align=al)
            _set_cell_border(c, **_thin("CBD5E1")) # lighter border for data
            _set_cell_bg(c, bg_color)

    if col_widths:
        for ci, w in enumerate(col_widths[:n_cols]):
            _col_w(tbl, ci, w)
    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# Chart generation helpers  (return BytesIO PNG)
# ─────────────────────────────────────────────────────────────────────────────

def _chart_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf


def _chart_element_profile(elements, title, y_key, y_label, color=CHART_BRAND,
                            figsize=(5.5, 2.6)):
    """Line chart of a per-element metric across all elements."""
    xs   = [f"S{e['stage']}-E{e['position']}" for e in elements]
    ys   = [e[y_key] for e in elements]
    fig, ax = plt.subplots(figsize=figsize)
    ax.plot(xs, ys, marker='o', color=color, linewidth=1.8, markersize=4)
    ax.fill_between(range(len(ys)), ys, alpha=0.12, color=color)
    ax.set_title(title); ax.set_ylabel(y_label)
    ax.set_xticks(range(len(xs))); ax.set_xticklabels(xs, rotation=45, ha='right', fontsize=7)
    ax.yaxis.set_major_formatter(mticker.FormatStrFormatter('%.1f'))
    # Stage separators
    stages = [e['stage'] for e in elements]
    for i in range(1, len(stages)):
        if stages[i] != stages[i-1]:
            ax.axvline(x=i-0.5, color='gray', linestyle='--', linewidth=0.7, alpha=0.6)
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_dual_profile(elements, title, y1_key, y2_key, y1_label, y2_label,
                         figsize=(5.5, 2.6)):
    """Dual-axis line chart (e.g., flux + NDP)."""
    xs  = [f"S{e['stage']}-E{e['position']}" for e in elements]
    y1s = [e[y1_key] for e in elements]
    y2s = [e[y2_key] for e in elements]
    fig, ax1 = plt.subplots(figsize=figsize)
    ax2 = ax1.twinx()
    ax1.plot(xs, y1s, marker='o', color=CHART_BRAND, linewidth=1.8, markersize=4, label=y1_label)
    ax2.plot(xs, y2s, marker='s', color=CHART_ORANGE,  linewidth=1.8, markersize=4, label=y2_label)
    ax1.set_title(title); ax1.set_ylabel(y1_label, color=CHART_BRAND)
    ax2.set_ylabel(y2_label, color=CHART_ORANGE)
    ax1.set_xticks(range(len(xs))); ax1.set_xticklabels(xs, rotation=45, ha='right', fontsize=7)
    lines  = ax1.get_lines() + ax2.get_lines()
    labels = [y1_label, y2_label]
    ax1.legend(lines, labels, fontsize=7, loc='upper right')
    ax1.spines['top'].set_visible(False)
    fig.tight_layout()
    return _chart_bytes(fig)

def _chart_dual_profile_manual(elements, y1s, y2s, title, y1_label, y2_label, figsize=(5.5, 2.6)):
    """Dual-axis line chart with manually scaled values."""
    xs  = [f"S{e['stage']}-E{e['position']}" for e in elements]
    fig, ax1 = plt.subplots(figsize=figsize)
    ax2 = ax1.twinx()
    ax1.plot(xs, y1s, marker='o', color=CHART_BRAND, linewidth=1.8, markersize=4, label=y1_label)
    ax2.plot(xs, y2s, marker='s', color=CHART_ORANGE,  linewidth=1.8, markersize=4, label=y2_label)
    ax1.set_title(title); ax1.set_ylabel(y1_label, color=CHART_BRAND)
    ax2.set_ylabel(y2_label, color=CHART_ORANGE)
    ax1.set_xticks(range(len(xs))); ax1.set_xticklabels(xs, rotation=45, ha='right', fontsize=7)
    lines  = ax1.get_lines() + ax2.get_lines()
    labels = [y1_label, y2_label]
    ax1.legend(lines, labels, fontsize=7, loc='upper right')
    ax1.spines['top'].set_visible(False)
    fig.tight_layout()
    return _chart_bytes(fig)

def _chart_element_profile_manual(elements, ys, title, y_label, color, figsize=(3.0, 2.5)):
    """Single-axis line chart with manually scaled values."""
    xs  = [f"S{e['stage']}-E{e['position']}" for e in elements]
    fig, ax = plt.subplots(figsize=figsize)
    ax.plot(xs, ys, marker='o', color=color, linewidth=1.8, markersize=4)
    ax.set_title(title, fontsize=8)
    ax.set_ylabel(y_label, fontsize=7, color=color)
    ax.set_xticks(range(len(xs))); ax.set_xticklabels(xs, rotation=45, ha='right', fontsize=6)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.tick_params(axis='both', which='major', labelsize=6)
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_bar_flow(stages, figsize=(5.5, 2.6)):
    """Grouped bar chart: Feed / Permeate / Concentrate per stage."""
    snames = [f"Stage {s['stage']}" for s in stages]
    feed   = [s['feed_flow']  for s in stages]
    perm   = [s['perm_flow']  for s in stages]
    conc   = [s['conc_flow']  for s in stages]
    x = np.arange(len(snames)); w = 0.25
    fig, ax = plt.subplots(figsize=figsize)
    ax.bar(x - w,   feed, w, label='Feed',        color=CHART_BLUE,   alpha=0.85)
    ax.bar(x,       perm, w, label='Permeate',    color=CHART_BRAND,  alpha=0.85)
    ax.bar(x + w,   conc, w, label='Concentrate', color=CHART_ORANGE, alpha=0.85)
    ax.set_title('Stage Flow Distribution (m³/h)')
    ax.set_ylabel('Flow (m³/h)')
    ax.set_xticks(x); ax.set_xticklabels(snames)
    ax.legend(fontsize=7)
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_ion_rejection(summary, figsize=(5.5, 2.8)):
    """Horizontal bar chart of % rejection per ion."""
    feed_ions = summary.get('feed_tds_ions', {})
    perm_ions = summary.get('perm_ions', {})
    # Build rejection from perm_ions vs feed_ions (use feed_water passed separately)
    ions, rejs = [], []
    for ion, fc in feed_ions.items():
        if fc > 0:
            pc  = perm_ions.get(ion, 0)
            rej = max(0, (1 - pc / fc) * 100)
            ions.append(ion)
            rejs.append(rej)
    if not ions:
        return None
    colors = [CHART_BRAND if r >= 99 else (CHART_BLUE if r >= 95 else CHART_ORANGE) for r in rejs]
    fig, ax = plt.subplots(figsize=figsize)
    bars = ax.barh(ions, rejs, color=colors, alpha=0.85)
    ax.set_xlim(0, 105)
    ax.set_title('Ion Rejection (%)'); ax.set_xlabel('Rejection (%)')
    for bar, rej in zip(bars, rejs):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2,
                f'{rej:.1f}%', va='center', fontsize=7)
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_ion_bar_comparison(feed_w, perm_ions, conc_ions, figsize=(6.0, 3.0)):
    """Grouped bar: Feed vs Permeate vs Concentrate for active ions."""
    active = {ion: val for ion, val in feed_w.items()
              if isinstance(val, (int, float)) and val > 0
              and ion not in ('temperature', 'ph', 'turbidity', 'tss', 'tds')}
    if not active:
        return None
    # Map feed_water keys to ion abbreviations
    key_map = {
        'sodium': 'Na', 'calcium': 'Ca', 'magnesium': 'Mg', 'potassium': 'K',
        'chloride': 'Cl', 'sulfate': 'SO4', 'bicarbonate': 'HCO3',
        'barium': 'Ba', 'strontium': 'Sr', 'fluoride': 'F',
        'silica': 'SiO2', 'boron': 'B', 'nitrate': 'NO3', 'phosphate': 'PO4',
        'ammonium': 'NH4',
    }
    ions  = [key_map.get(k, k) for k in active]
    feeds = list(active.values())
    perms = [perm_ions.get(ion, 0) for ion in ions]
    concs = [conc_ions.get(ion, 0) for ion in ions]

    x = np.arange(len(ions)); w = 0.28
    fig, ax = plt.subplots(figsize=figsize)
    ax.bar(x - w,   feeds, w, label='Feed',        color=CHART_BLUE,   alpha=0.85)
    ax.bar(x,       perms, w, label='Permeate',    color=CHART_BRAND,  alpha=0.85)
    ax.bar(x + w,   concs, w, label='Concentrate', color=CHART_ORANGE, alpha=0.85)
    ax.set_title('Ion Concentrations: Feed vs Permeate vs Concentrate (mg/L)')
    ax.set_ylabel('Concentration (mg/L)')
    ax.set_xticks(x); ax.set_xticklabels(ions, rotation=45, ha='right', fontsize=7)
    ax.legend(fontsize=7)
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_tds_profile(elements, figsize=(5.5, 2.6)):
    """Show feed-side and concentrate-side TDS rising along element train."""
    xs  = [f"S{e['stage']}-E{e['position']}" for e in elements]
    tds = [sum(v for k,v in e.get('feed_ions', {}).items() if k not in ("SiO2", "B", "CO2")) for e in elements]
    # Use rej_ions (concentrate-side ions) or conc_ions if present
    tds_conc = [
        sum(v for k,v in e.get('conc_ions', e.get('rej_ions', e.get('feed_ions', {}))).items() if k not in ("SiO2", "B", "CO2"))
        for e in elements
    ]
    fig, ax = plt.subplots(figsize=figsize)
    ax.plot(xs, tds,      marker='o', color=CHART_BLUE,   linewidth=1.8, markersize=4, label='Feed TDS')
    ax.plot(xs, tds_conc, marker='s', color=CHART_ORANGE, linewidth=1.8, markersize=4, label='Conc TDS')
    ax.fill_between(range(len(tds_conc)), tds, tds_conc, alpha=0.08, color=CHART_ORANGE)
    ax.set_title('TDS Profile Along Element Train (mg/L)')
    ax.set_ylabel('TDS (mg/L)')
    ax.set_xticks(range(len(xs))); ax.set_xticklabels(xs, rotation=45, ha='right', fontsize=7)
    ax.legend(fontsize=7)
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_capex_pie(capex, figsize=(4.0, 3.0)):
    labels = ['Membranes', 'Vessels', 'HP Pump', 'Booster', 'Install+Cont.']
    vals   = [
        capex.get('membranes_inr', 0),
        capex.get('vessels_inr', 0),
        capex.get('hp_pump_inr', 0),
        capex.get('booster_pump_inr', 0),
        capex.get('ic_inr', 0) + capex.get('contingency_inr', 0),
    ]
    colors = [CHART_BRAND, CHART_BLUE, CHART_ORANGE, CHART_GRAY, '#9C27B0']
    fig, ax = plt.subplots(figsize=figsize)
    wedges, texts, autotexts = ax.pie(
        vals, labels=labels, colors=colors,
        autopct='%1.1f%%', startangle=140,
        textprops={'fontsize': 7}, pctdistance=0.75
    )
    for t in autotexts: t.set_fontsize(6.5)
    ax.set_title('CAPEX Breakdown')
    fig.tight_layout()
    return _chart_bytes(fig)


def _chart_opex_bar(opex, figsize=(4.0, 2.6)):
    labels = ['Energy', 'Membrane Repl.']
    vals   = [opex.get('energy_cost_pa_inr', 0), opex.get('membrane_repl_pa_inr', 0)]
    colors = [CHART_ORANGE, CHART_BLUE]
    fig, ax = plt.subplots(figsize=figsize)
    bars = ax.bar(labels, [v/1e5 for v in vals], color=colors, alpha=0.85, width=0.45)
    ax.set_title('Annual OPEX Breakdown')
    ax.set_ylabel('Cost (₹ Lakhs/year)')
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.02,
                f'₹{v/1e5:.2f}L', ha='center', fontsize=7)
    fig.tight_layout()
    return _chart_bytes(fig)


def _insert_chart(doc, buf, width_in=6.0):
    """Insert a chart image into the document."""
    if buf is None:
        return
    doc.add_picture(buf, width=Inches(width_in))


def _insert_two_charts(doc, buf1, buf2, width_in=3.0):
    """Place two chart images side-by-side in a borderless table."""
    if buf1 is None and buf2 is None:
        return
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, buf in enumerate([buf1, buf2]):
        cell = tbl.rows[0].cells[ci]
        _set_cell_border(cell, **_no_border())
        if buf:
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(buf, width=Inches(width_in))


# ─────────────────────────────────────────────────────────────────────────────
# Header / footer helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_header_block(doc, proj_name):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    left = tbl.rows[0].cells[0]
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Try inserting logo, fallback to text
    logo_path = os.path.join(os.path.dirname(__file__), 'assets', 'permionics_logo.png')
    if os.path.exists(logo_path):
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.8))
    else:
        _cell_text(left, 'PERMIONICS MEMBRANES PVT. LTD.', bold=True, size=13, color=BRAND_COLOR)

    right = tbl.rows[0].cells[1]
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = right.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, 'PACE: PERMIONICS ANALYTICAL CALCULATION ENGINE\n', bold=True, size=9, color=BRAND_COLOR)
    _run(p, 'Customized Membrane Solutions', size=8, color=BLACK)
    for c in [left, right]:
        _set_cell_border(c, **_no_border())
    _hr(doc)


def _add_footer(doc, proj_name, date_str):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        ft = footer.paragraphs[0]
        ft.clear()
        pPr  = ft._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top  = OxmlElement('w:top')
        top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1'); top.set(qn('w:color'), '4169E1')
        pBdr.append(top); pPr.append(pBdr)
        _run(ft, f'Project: {proj_name}     Created: {date_str}     Page ', size=8, color=BRAND_COLOR)
        # Page number field
        for tag, text in [('begin', ''), ('', 'PAGE'), ('end', '')]:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), tag if tag else 'separate')
            it = OxmlElement('w:instrText'); it.text = text
            r  = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            c = OxmlElement('w:color'); c.set(qn('w:val'), '4169E1')
            rPr.append(c); r.append(rPr)
            r.append(fc if tag else it)
            ft._p.append(r)


# ─────────────────────────────────────────────────────────────────────────────
# Report Generator class
# ─────────────────────────────────────────────────────────────────────────────

class ReportGenerator:
    def __init__(self):
        self.today = datetime.now().strftime('%d/%m/%Y')

    def _chart_pfd_diagram(self, ro_results, sr):
        if not ro_results:
            return None
        
        stages = ro_results.get('stages', [])
        if not stages:
            return None
            
        summary = ro_results.get('summary', {})
        boosters = ro_results.get('booster_pumps', [])
        
        # Parse vessel array
        vps = summary.get('vessels_per_stage', [4, 2])
        n_el = sr.get('elements_per_vessel', summary.get('elements_per_vessel', 6))
        
        # Determine number of stages
        num_stages = len(vps)
        rEdge = 4.5 + num_stages * 3.0
        
        fig, ax = plt.subplots(figsize=(6.5, 3.2), dpi=150)
        ax.axis('off')
        ax.set_xlim(-1, rEdge)
        ax.set_ylim(-2, 6)
        
        # Helper to draw text
        def txt(x, y, s, color='#1e293b', size=7.5, weight='normal', ha='center', zorder=3):
            ax.text(x, y, s, color=color, fontsize=size, fontweight=weight, ha=ha, va='center', fontfamily='DejaVu Sans', zorder=zorder)

        # Helper to draw rect
        from matplotlib.patches import FancyBboxPatch, Circle, Polygon, Ellipse
        def draw_rect(x, y, w, h, fill, stroke, rx=0.1, zorder=2):
            p = FancyBboxPatch((x, y), w, h, boxstyle=f"round,pad=0,rounding_size={rx}", facecolor=fill, edgecolor=stroke, linewidth=1.2, zorder=zorder)
            ax.add_patch(p)
            
        # Helper to draw lines
        def draw_line(x1, y1, x2, y2, color, lw=1.5, style='-', zorder=1):
            ax.plot([x1, x2], [y1, y2], color=color, linewidth=lw, linestyle=style, zorder=zorder)

        # Draw PFD Header background
        draw_rect(-0.8, 5.0, rEdge + 1.2, 0.8, '#f1f5f9', '#cbd5e1', rx=0.05, zorder=2)
        txt(-0.4 + rEdge/2, 5.4, 'Process Flow Diagram', color='#1e293b', size=10, weight='bold', zorder=3)
        
        # Coordinate definitions
        N_ML = 1.0
        N_PP = 1.2
        N_VW = 1.4
        N_VH = 0.5
        N_VGY = 0.25
        
        feedX = 0.2
        filtX = feedX + N_PP
        hpX = filtX + N_PP
        p1CY = 2.0
        
        # 1. Feed Water
        # Draw feed water triangle pointing to the right
        feed_poly = Polygon([[feedX, p1CY - 0.4], [feedX, p1CY + 0.4], [feedX + 0.6, p1CY]], closed=True, facecolor='#eff6ff', edgecolor='#2563eb', linewidth=1.5, zorder=3)
        ax.add_patch(feed_poly)
        txt(feedX + 0.3, p1CY - 0.8, 'Feed', weight='bold', size=7, zorder=3)
        feed_flow = summary.get('feed_flow', 0)
        if feed_flow > 0:
            txt(feedX + 0.3, p1CY + 0.7, f"{feed_flow:.1f} m³/h", color='#2563eb', weight='bold', size=7.5, zorder=3)
            
        # 2. Pre-Filter
        draw_rect(filtX, p1CY - 0.5, 0.5, 1.0, '#f8fafc', '#cbd5e1', rx=0.05, zorder=2)
        # filter lines
        for y in [p1CY - 0.25, p1CY, p1CY + 0.25]:
            draw_line(filtX + 0.05, y, filtX + 0.45, y, '#cbd5e1', lw=1.0, zorder=3)
        # Filter top lid ellipse
        filter_ellipse = Ellipse((filtX + 0.25, p1CY + 0.5), 0.5, 0.1, facecolor='#e2e8f0', edgecolor='#cbd5e1', linewidth=1.2, zorder=3)
        ax.add_patch(filter_ellipse)
        txt(filtX + 0.25, p1CY - 0.8, 'Pre-Filter', size=7, zorder=3)
        
        # Connection Feed -> Filter -> Pump
        draw_line(feedX + 0.6, p1CY, filtX, p1CY, '#3b82f6', lw=2.0, zorder=1)
        draw_line(filtX + 0.5, p1CY, hpX, p1CY, '#3b82f6', lw=2.0, zorder=1)
        
        # 3. HP Pump
        pump_circle = Circle((hpX + 0.3, p1CY), 0.35, facecolor='#eff6ff', edgecolor='#2563eb', linewidth=2.0, zorder=2)
        ax.add_patch(pump_circle)
        pump_tri = Polygon([[hpX + 0.15, p1CY - 0.2], [hpX + 0.15, p1CY + 0.2], [hpX + 0.42, p1CY]], closed=True, facecolor='#2563eb', zorder=3)
        ax.add_patch(pump_tri)
        txt(hpX + 0.3, p1CY - 0.8, 'HP Pump', weight='bold', size=7, zorder=3)
        feed_press = summary.get('feed_pressure_bar', 0)
        if feed_press > 0:
            txt(hpX + 0.3, p1CY + 0.7, f"{feed_press:.1f} bar", color='#2563eb', weight='bold', size=7.5, zorder=3)
            
        # Connection Pump -> Stage 1
        sx0 = hpX + 1.2
        draw_line(hpX + 0.65, p1CY, sx0, p1CY, '#3b82f6', lw=2.0, zorder=1)
        
        # 4. RO Stages
        stage_x = []
        cur_x = sx0
        booster_map = {bp.get('from_stage'): bp for bp in boosters if bp.get('required')}
        
        # Helper function to size block height
        def stH(n):
            return n * N_VH + (n - 1) * N_VGY
            
        for si, nv in enumerate(vps):
            stage_x.append(cur_x)
            cur_x += N_VW + 1.2
            if booster_map.get(si + 1):
                cur_x += 0.5 # Add space for booster pump
                
        permY = -0.8
        
        for si, nv in enumerate(vps):
            cx = stage_x[si]
            sH = stH(nv)
            topY = p1CY - sH / 2
            
            # Stage Header Label
            draw_rect(cx + N_VW/2 - 0.6, topY + sH + 0.4, 1.2, 0.35, '#2563eb', '#2563eb', rx=0.03, zorder=3)
            txt(cx + N_VW/2, topY + sH + 0.55, f"Stage {si+1}", color='#ffffff', weight='bold', size=7, zorder=4)
            
            # Stage recovery metrics
            stage_data = stages[si] if si < len(stages) else {}
            rec_pct = stage_data.get('recovery', 0) * 100
            draw_rect(cx + N_VW/2 - 0.7, topY + sH + 0.02, 1.4, 0.35, '#f8fafc', '#cbd5e1', rx=0.03, zorder=3)
            txt(cx + N_VW/2, topY + sH + 0.17, f"Rec: {rec_pct:.1f}%", color='#475569', size=6.5, zorder=4)
            
            # Draw vertical manifolds
            # Feed manifold (left)
            draw_line(cx - 0.1, topY, cx - 0.1, topY + sH - N_VH/2, '#3b82f6', lw=3.5, zorder=2)
            # Concentrate manifold (right)
            draw_line(cx + N_VW + 0.1, topY, cx + N_VW + 0.1, topY + sH - N_VH/2, '#ef4444', lw=3.5, zorder=2)
            
            # Draw vessels
            for vi in range(nv):
                vy = topY + vi * (N_VH + N_VGY)
                # draw vessel box (zorder=2)
                draw_rect(cx, vy, N_VW, N_VH, '#f0f7ff', '#3b82f6', rx=0.05, zorder=2)
                # vessel stubs
                draw_line(cx - 0.1, vy + N_VH/2, cx, vy + N_VH/2, '#3b82f6', lw=1.2, zorder=2)
                draw_line(cx + N_VW, vy + N_VH/2, cx + N_VW + 0.1, vy + N_VH/2, '#ef4444', lw=1.2, zorder=2)
                
                # Element dividers inside vessel
                ew = (N_VW - 0.15) / n_el
                for ei in range(n_el - 1):
                    draw_line(cx + 0.075 + ew*(ei+1), vy + 0.05, cx + 0.075 + ew*(ei+1), vy + N_VH - 0.05, '#bfdbfe', lw=0.8, style='--', zorder=3)
                    
                # vessel text
                txt(cx + N_VW/2, vy + N_VH/2, f"{si+1}-{vi+1}", color='#1e293b', size=6.5, zorder=3)
                
                # permeate line down (low zorder=1 to pass UNDER lower vessels)
                draw_line(cx + N_VW/2, vy, cx + N_VW/2, permY, '#10b981', lw=1.0, style='--', zorder=1)
                
            # Connect input line to feed manifold
            draw_line(cx - 0.5, p1CY, cx - 0.1, p1CY, '#3b82f6', lw=2.0, zorder=1)
            
            # Connect concentrate manifold to next stage or output
            if si < len(vps) - 1:
                next_cx = stage_x[si+1]
                from_x = cx + N_VW + 0.1
                to_x = next_cx - 0.5
                
                bp = booster_map.get(si + 1)
                if bp:
                    mid_x = (from_x + to_x) / 2
                    draw_line(from_x, p1CY, mid_x - 0.25, p1CY, '#ef4444', lw=2.0, zorder=1)
                    draw_line(mid_x + 0.25, p1CY, to_x, p1CY, '#3b82f6', lw=2.0, zorder=1)
                    
                    # Booster pump circle
                    bp_circle = Circle((mid_x, p1CY), 0.28, facecolor='#fffbeb', edgecolor='#d97706', linewidth=1.5, zorder=2)
                    ax.add_patch(bp_circle)
                    bp_tri = Polygon([[mid_x - 0.1, p1CY - 0.15], [mid_x - 0.1, p1CY + 0.15], [mid_x + 0.12, p1CY]], closed=True, facecolor='#d97706', zorder=3)
                    ax.add_patch(bp_tri)
                    
                    txt(mid_x, p1CY - 0.55, 'Booster', weight='bold', size=6, zorder=3)
                    txt(mid_x, p1CY + 0.45, f"+{bp.get('boost_dp_bar', 0):.1f} bar", color='#d97706', weight='bold', size=6, zorder=3)
                else:
                    draw_line(from_x, p1CY, to_x, p1CY, '#ef4444', lw=2.0, zorder=1)
            else:
                # Final reject stream
                from_x = cx + N_VW + 0.1
                to_x = from_x + 1.0
                draw_line(from_x, p1CY, to_x, p1CY, '#ef4444', lw=2.5, zorder=1)
                ax.annotate('', xy=(to_x, p1CY), xytext=(to_x - 0.25, p1CY), arrowprops=dict(arrowstyle="->", color='#ef4444', lw=2.5), zorder=3)
                txt(to_x, p1CY - 0.4, 'Reject', color='#ef4444', weight='bold', size=7, zorder=3)
                conc_flow = summary.get('conc_flow', 0)
                if conc_flow > 0:
                    txt(to_x, p1CY + 0.4, f"{conc_flow:.1f} m³/h", color='#ef4444', weight='bold', size=7.5, zorder=3)

        # 5. Permeate Collection header line - extended to directly connect to the Product Tank
        p1PermLx = stage_x[0] + N_VW/2
        p1PermRx = stage_x[-1] + N_VW/2
        tank_x = p1PermRx + 1.8
        
        draw_line(p1PermLx, permY, tank_x - 0.3, permY, '#10b981', lw=3.0, zorder=1)
        ax.annotate('', xy=(tank_x - 0.3, permY), xytext=(tank_x - 0.7, permY), arrowprops=dict(arrowstyle="->", color='#10b981', lw=3.0), zorder=3)
        txt(p1PermRx + 0.6, permY - 0.4, 'Permeate Stream', color='#10b981', weight='bold', size=7, zorder=3)
        
        perm_flow = summary.get('perm_flow', 0)
        if perm_flow > 0:
            txt(p1PermRx + 0.6, permY + 0.3, f"{perm_flow:.1f} m³/h", color='#10b981', weight='bold', size=7.5, zorder=3)
            
        # 6. Permeate Tank (draw cylindrical tank with body and top/bottom ellipses)
        draw_rect(tank_x - 0.3, permY - 0.5, 0.6, 0.8, '#ecfdf5', '#10b981', rx=0.02, zorder=2)
        top_ellipse = Ellipse((tank_x, permY + 0.3), 0.6, 0.12, facecolor='#d1fae5', edgecolor='#10b981', linewidth=1.2, zorder=3)
        ax.add_patch(top_ellipse)
        bot_ellipse = Ellipse((tank_x, permY - 0.5), 0.6, 0.12, facecolor='#d1fae5', edgecolor='#10b981', linewidth=1.5, zorder=3)
        ax.add_patch(bot_ellipse)
        txt(tank_x, permY - 0.8, 'Product Tank', color='#047857', weight='bold', size=6.5, zorder=3)
        
        return _chart_bytes(fig)

    def _setup_doc(self):
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.8)
            s.left_margin = Cm(1.8); s.right_margin = Cm(1.8)
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(9.5)
        return doc

    # ── Page 1: Cover + System Overview ─────────────────────────────────────

    def _render_pfd_svg(self, svg_str):
        import tempfile, os, io
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        try:
            with tempfile.NamedTemporaryFile(suffix=".svg", delete=False) as f:
                f.write(svg_str.encode('utf-8'))
                temp_name = f.name
            drawing = svg2rlg(temp_name)
            buf = io.BytesIO()
            # Scaling up for better resolution in report
            drawing.renderScale = 2.0 
            renderPM.drawToFile(drawing, buf, fmt="PNG", dpi=300)
            os.remove(temp_name)
            buf.seek(0)
            return buf
        except Exception as e:
            print(f"SVG Render error: {e}")
            return None

    def _page_uf_overview(self, doc, sr):
        uf_results = sr.get('uf_results')
        if not uf_results: return

        _heading(doc, 'UF Summary Report', size=16)

        overview = uf_results.get('overview', {})
        op_cond = uf_results.get('operating_conditions', {})
        
        ud = sr.get('units', {})
        flow_lbl = lbl_flow(ud)
        press_lbl = lbl_press(ud)
        flux_lbl = lbl_flux(ud)

        _heading(doc, 'UF System Overview', size=11)
        _prop_table(doc, [
            ('UF Module Model',             overview.get('module_type', '—')),
            ('Total Modules',               str(overview.get('total_modules', '—'))),
            ('Gross Feed Flow',             f"{conv_flow(overview.get('gross_feed_m3h', 0), ud):.1f} {flow_lbl}"),
            ('Net Product Flow',            f"{conv_flow(overview.get('net_product_m3h', 0), ud):.1f} {flow_lbl}"),
            ('Recovery',                    f"{overview.get('recovery_pct', 0):.1f} %"),
            ('Filtration Flux',             f"{conv_flux(op_cond.get('filtration_flux_lmh', 0), ud):.1f} {flux_lbl}"),
            ('Design TMP',                  f"{conv_press(overview.get('tmp_design_bar', 0), ud):.2f} {press_lbl}"),
            ('Filtration Duration',         f"{op_cond.get('filtration_duration_min', 0):.1f} min"),
            ('Backwash Duration',           f"{op_cond.get('backwash_duration_min', 0):.1f} min"),
        ], widths=(7.5, 6.5))
        _spacer(doc)
        
        doc.add_page_break()

    def _page_system_overview(self, doc, sr):
        pd = sr.get('project_details', {})
        proj_name = pd.get('name') or sr.get('project_name', 'PACE Report')
        _add_header_block(doc, proj_name)
        _heading(doc, 'RO Summary Report', size=16)

        # ── Project Information Table ──
        if pd:
            _heading(doc, 'Project Information', size=11)
            
            # Combine names if present
            full_name = f"{pd.get('firstName', '')} {pd.get('lastName', '')}".strip()
            
            pd_rows = [
                ['Project Name', proj_name],
                ['Prepared By', full_name or '—'],
                ['Company', pd.get('company') or '—'],
                ['Email', pd.get('email') or '—'],
                ['Office', pd.get('office') or '—'],
                ['Mobile', pd.get('mobile') or '—'],
                ['Address', f"{pd.get('street', '')}, {pd.get('city', '')}, {pd.get('country', '')}".strip(", ") or '—'],
            ]
            
            _prop_table(doc, pd_rows)
            _spacer(doc)

        _heading(doc, 'RO System Flow Diagram', size=11)

        pfd_buf = None
        if sr.get('pfd_png'):
            import base64
            import io
            try:
                # Strip data URL prefix if present (e.g. data:image/png;base64,...)
                b64_data = sr['pfd_png'].split(",")[1] if "," in sr['pfd_png'] else sr['pfd_png']
                image_data = base64.b64decode(b64_data)
                pfd_buf = io.BytesIO(image_data)
            except Exception as e:
                print(f"PNG Decode error: {e}")
        elif sr.get('pfd_svg'):
            pfd_buf = self._render_pfd_svg(sr.get('pfd_svg'))
        
        if not pfd_buf:
            ro_main = sr.get('pass1_results') or sr.get('ro_results')
            pfd_buf = self._chart_pfd_diagram(ro_main, sr)
        if pfd_buf:
            _insert_chart(doc, pfd_buf, width_in=6.2)
        else:
            train = sr.get('technology_train', '1P-RO')
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            diagram = (
                "             [4] Concentrate\n"
                "              ↑\n"
                " [1] Feed ──► [2] Pump ──► [Pass 1] ─────────────────────────────► [6] Product\n"
            ) if '1P' in train else (
                " [1] Feed ──► [Pass 1] ──► [Perm1] ──► [Pass 2] ──► [6] Product\n"
                "                  ↓                          ↓\n"
                "               Conc 1                    Conc 2\n"
            )
            _run(p, diagram)
            p.runs[-1].font.name = 'Courier New'; p.runs[-1].font.size = Pt(8)

        # ── Stream flow table ──
        ro_main = sr.get('pass1_results') or sr.get('ro_results') or {}
        sm  = ro_main.get('summary', {})
        rec = sm.get('total_recovery', 0) * 100
        ud = sr.get('units', {})
        flow_lbl = lbl_flow(ud)
        press_lbl = lbl_press(ud)
        flux_lbl = lbl_flux(ud)

        _heading(doc, 'RO Flow Table (Stream Level)')
        _data_table(doc,
            [['#', 'Description', f'Flow ({flow_lbl})', 'TDS (mg/L)', f'Pressure ({press_lbl})']],
            [
                ['1', 'Raw Feed to RO System',           f"{conv_flow(sr.get('target_flow_m3h', sm.get('feed_flow', 0)), ud):.1f}",  f"{sm.get('feed_tds', 0):.1f}", '0.0'],
                ['2', 'Net Feed to Pass 1',               f"{conv_flow(sm.get('feed_flow', 0), ud):.1f}",  f"{sm.get('feed_tds', 0):.1f}", f"{conv_press(sm.get('feed_pressure_bar', 0), ud):.1f}"],
                ['4', 'Total Concentrate from Pass 1',    f"{conv_flow(sm.get('conc_flow', 0), ud):.1f}",  f"{sm.get('conc_tds', 0):.0f}", f"{conv_press(sm.get('conc_pressure_bar', 0), ud):.1f}"],
                ['6', 'Net Product from RO System',       f"{conv_flow(sm.get('perm_flow', 0), ud):.1f}",  f"{sm.get('perm_tds', 0):.2f}", '0.0'],
            ],
            col_widths=[0.7, 6.5, 2.5, 2.5, 2.2]
        )
        _spacer(doc)

        # ── System overview table ──
        vps  = sm.get('vessels_per_stage', [1])
        epv  = sr.get('elements_per_vessel', sm.get('elements_per_vessel', 6))
        n_el = sum(vps) * epv
        _heading(doc, 'RO System Overview')

        # Top row: units / recovery summary
        tbl_top = doc.add_table(rows=2, cols=6)
        top_data = [
            ['Total # of Units', '1', 'Online =', '1', 'Standby =', '0'],
            ['RO Recovery', f'{rec:.1f} %', 'System Flow Rate', f"{conv_flow(sr.get('target_flow_m3h', 0), ud):.1f} {flow_lbl}",
             'Net Product', f"{conv_flow(sm.get('perm_flow', 0), ud):.1f} {flow_lbl}"],
        ]
        for ri, row in enumerate(top_data):
            for ci, val in enumerate(row):
                c = tbl_top.rows[ri].cells[ci]
                _cell_text(c, val, bold=(ci % 2 == 0), size=9)
                _set_cell_border(c, **_thin())
        _spacer(doc, 4)

        # Detail property table
        _prop_table(doc, [
            ('Pass',                       'Pass 1'),
            ('Membrane Model',             sr.get('ro_membrane', '—')),
            ('Number of Elements',         str(n_el)),
            ('Vessels per Stage',          ' | '.join(str(v) for v in vps)),
            ('Elements per Vessel',        str(sr.get('elements_per_vessel', 6))),
            ('Feed Flow per Pass',         f"{conv_flow(sm.get('feed_flow', 0), ud):.1f} {flow_lbl}"),
            ('Feed TDS',                   f"{sm.get('feed_tds', 0):.1f} mg/L"),
            ('Feed Pressure',              f"{conv_press(sm.get('feed_pressure_bar', 0), ud):.1f} {press_lbl}"),
            ('Permeate Flow per Pass',     f"{conv_flow(sm.get('perm_flow', 0), ud):.1f} {flow_lbl}"),
            ('Concentrate Flow',           f"{conv_flow(sm.get('conc_flow', 0), ud):.1f} {flow_lbl}"),
            ('Average Flux',               f"{conv_flux(sm.get('avg_flux_lmh', 0), ud):.1f} {flux_lbl}"),
            ('Permeate TDS',               f"{sm.get('perm_tds', 0):.2f} mg/L"),
            ('Concentrate TDS',            f"{sm.get('conc_tds', 0):.0f} mg/L"),
            ('System Recovery',            f"{rec:.1f} %"),
            ('HP Pump Power',              f"{sm.get('hp_pump_power_kw', 0):.2f} kW"),
            ('Booster Pump Power',         f"{sm.get('booster_pump_power_kw', 0):.2f} kW"),
            ('Total Power',                f"{sm.get('total_power_kw', 0):.2f} kW"),
            ('Specific Energy Consumption',f"{sm.get('sec_kwh_m3', 0):.3f} kWh/m³"),
            ('Temperature',                f"{sr.get('feed_water_used', {}).get('temperature', 25):.1f} °C"),
            ('Feed pH',                    f"{sr.get('feed_water_used', {}).get('ph', 7):.1f}"),
        ], widths=(7.5, 6.5))

        self._footnote(doc)

    # ── Page 2: Stage + Element flow tables ─────────────────────────────────

    def _page_flow_tables(self, doc, ro: dict, membrane: str, pass_label: str, sr: dict):
        sm       = ro.get('summary', {})
        stages   = ro.get('stages', [])
        elements = ro.get('elements', [])
        vps      = sm.get('vessels_per_stage', [1])
        
        ud = sr.get('units', {})
        flow_lbl = lbl_flow(ud)
        press_lbl = lbl_press(ud)
        flux_lbl = lbl_flux(ud)

        # Stage-level table
        _heading(doc, f'RO Flow Table (Stage Level) – {pass_label}')
        stage_rows = []
        for i, stg in enumerate(stages):
            sn   = stg['stage']
            vs   = vps[i] if i < len(vps) else 1
            s_el = sorted([e for e in elements if e['stage'] == sn], key=lambda x: x['position'])
            if not s_el: continue
            avg_flux  = sum(e.get('flux', 0) for e in s_el) / len(s_el)
            avg_ndp   = sum(e.get('ndp', 0)  for e in s_el) / len(s_el)
            fp        = s_el[0].get('feed_pressure', 0)
            cp        = s_el[-1].get('conc_pressure', 0)
            dp        = fp - cp
            epv       = len(s_el)
            stage_rows.append([
                str(sn), membrane, str(vs), str(epv),
                f"{conv_flow(stg['feed_flow'], ud):.1f}",  f"{conv_press(fp, ud):.1f}",
                f"{conv_flow(stg['perm_flow'], ud):.1f}",  f"{conv_flow(stg['conc_flow'], ud):.1f}",
                f"{conv_press(cp, ud):.1f}", f"{conv_press(dp, ud):.1f}", f"{conv_flux(avg_flux, ud):.1f}", f"{conv_press(avg_ndp, ud):.1f}",
                f"{stg['recovery']*100:.1f}",
            ])
        _data_table(doc,
            [['Stage', 'Element Name', '#PV', 'Els/PV',
              f'Feed\nFlow\n({flow_lbl})', f'Feed\nPress\n({press_lbl})',
              f'Perm\nFlow\n({flow_lbl})', f'Conc\nFlow\n({flow_lbl})',
              f'Conc\nPress\n({press_lbl})', f'Press\nDrop\n({press_lbl})',
              f'Avg\nFlux\n({flux_lbl})', f'Avg\nNDP\n({press_lbl})', 'Recovery\n(%)']],
            stage_rows,
            col_widths=[0.8, 2.6, 0.7, 0.8, 1.3, 1.2, 1.3, 1.3, 1.2, 1.1, 1.1, 1.1, 1.3]
        )
        _spacer(doc, 8)

        # Element-level table
        _heading(doc, f'RO Flow Table (Element Level) – {pass_label}')
        el_rows = []
        for e in elements:
            feed_tds = sum(v for k,v in e.get('feed_ions', {}).items() if k not in ("SiO2", "B", "CO2"))
            perm_tds = sum(v for k,v in e.get('perm_ions', {}).items() if k not in ("SiO2", "B", "CO2"))
            el_rows.append([
                str(e.get('stage', '—')), str(e.get('position', '—')), membrane,
                f"{e.get('recovery', 0)*100:.1f}",
                f"{conv_flow(e.get('feed_flow', 0), ud):.1f}",  f"{conv_press(e.get('feed_pressure', 0), ud):.2f}",
                f"{feed_tds:.1f}",
                f"{conv_flow(e.get('perm_flow', 0), ud):.3f}",  f"{conv_flux(e.get('flux', 0), ud):.1f}",
                f"{conv_press(e.get('ndp', 0), ud):.2f}",         f"{e.get('beta', 0):.3f}",
                f"{perm_tds:.3f}",
            ])
        _data_table(doc,
            [['Stage', 'Elem', 'Model', 'Recovery\n(%)',
              f'Feed Flow\n({flow_lbl})', f'Feed Press\n({press_lbl})', 'Feed TDS\n(mg/L)',
              f'Perm Flow\n({flow_lbl})', f'Flux\n({flux_lbl})', f'NDP\n({press_lbl})',
              'β', 'Perm TDS\n(mg/L)']],
            el_rows,
            col_widths=[0.8, 0.7, 2.4, 1.1, 1.3, 1.3, 1.3, 1.3, 1.1, 1.0, 1.0, 1.3]
        )
        self._footnote(doc)
        _spacer(doc, 4)

        # (Stage Flow Bar Chart Removed as it's obvious from the table)
        
        _heading(doc, 'Hydraulic Performance & Visualization')
        
        # We need to scale chart arrays explicitly:
        flux_scaled = [conv_flux(e.get('flux', 0), ud) for e in elements]
        ndp_scaled = [conv_press(e.get('ndp', 0), ud) for e in elements]
        buf_flux_ndp = _chart_dual_profile_manual(elements, flux_scaled, ndp_scaled, 'Flux & NDP Along Element Train', f'Flux ({flux_lbl})', f'NDP ({press_lbl})')
        _insert_chart(doc, buf_flux_ndp, width_in=6.0)
        _spacer(doc, 4)

        press_scaled = [conv_press(e.get('feed_pressure', 0), ud) for e in elements]
        buf_press = _chart_element_profile_manual(elements, press_scaled, 'Feed Pressure Along Element Train', f'Pressure ({press_lbl})', CHART_ORANGE)
        buf_beta  = _chart_element_profile(elements, 'Concentration Polarization (β)', 'beta', 'β (–)', CHART_BLUE)
        _insert_two_charts(doc, buf_press, buf_beta, width_in=3.1)
        _spacer(doc, 4)
        
        # (TDS Profile Chart Removed as it's obvious from the table)
    # ── Page 3: Ion analysis ──────────────────────────────────────────────

    def _page_ion_analysis(self, doc, ro: dict, fw: dict):
        _heading(doc, 'RO Solute Concentrations & Rejection Analysis', size=16)
        sm        = ro.get('summary', {})
        perm_ions = sm.get('perm_ions', {})
        conc_ions = sm.get('conc_ions', {})

        # Ion table
        _heading(doc, 'Ion Concentration Table – Feed | Concentrate | Permeate | Rejection')
        ion_map = {
            'Na': 'sodium', 'Ca': 'calcium', 'Mg': 'magnesium', 'K': 'potassium',
            'Cl': 'chloride', 'SO4': 'sulfate', 'HCO3': 'bicarbonate',
            'Ba': 'barium', 'Sr': 'strontium', 'F': 'fluoride',
            'SiO2': 'silica', 'B': 'boron', 'NO3': 'nitrate',
            'PO4': 'phosphate', 'NH4': 'ammonium',
        }
        ion_rows = []
        active_ions = {}
        for ion, key in ion_map.items():
            fc = fw.get(key, 0.0)
            pc = perm_ions.get(ion, 0.0)
            cc = conc_ions.get(ion, 0.0)
            rej = (1 - pc / fc) * 100 if fc > 0 else 100.0
            ion_rows.append([ion, f'{fc:.2f}', f'{cc:.2f}', f'{pc:.4f}', f'{rej:.2f} %'])
            if fc > 0:
                active_ions[ion] = {'feed': fc, 'perm': pc}

        feed_tds = sm.get('feed_tds', 0)
        perm_tds = sm.get('perm_tds', 0)
        conc_tds = sm.get('conc_tds', 0)
        rej_tds  = (1 - perm_tds / feed_tds) * 100 if feed_tds > 0 else 100.0

        ion_rows.append(['TDS* (mg/L)',      f'{feed_tds:.1f}', f'{conc_tds:.1f}',  f'{perm_tds:.2f}',   f'{rej_tds:.2f} %'])
        ion_rows.append(['Conductivity (µS/cm)', f'{feed_tds*2.02:.0f}', f'{conc_tds*2.02:.0f}', f'{perm_tds*2.02:.1f}', '—'])

        _data_table(doc,
            [['Ion Species', 'Feed (mg/L)', 'Concentrate (mg/L)', 'Permeate (mg/L)', 'Rejection (%)']],
            ion_rows,
            col_widths=[3.0, 2.8, 3.2, 3.0, 2.8]
        )
        _spacer(doc, 8)
        self._footnote(doc)
        _spacer(doc, 4)

        # (Ion Rejection Bar Chart and Concentration Comparison Charts removed as they are obvious from the table)
        
    # ── Page 5: Warnings ─────────────────────────────────────────────────────

    def _page_warnings(self, doc, ro_results: dict, pass2_results: dict = None):
        _heading(doc, 'RO Design Warnings', size=16)

        all_warnings = []
        for w in ro_results.get('warnings', []):
            if 'Concentration Polarization' in w.get('type', ''):
                continue
            wc = dict(w); wc.setdefault('pass', '1'); all_warnings.append(wc)
        if pass2_results:
            for w in pass2_results.get('warnings', []):
                if 'Concentration Polarization' in w.get('type', ''):
                    continue
                wc = dict(w); wc['pass'] = '2'; all_warnings.append(wc)

        if all_warnings:
            rows = [[w.get('type', '—'), str(w.get('limit', '—')),
                     f"{w.get('value', 0):.3f}" if isinstance(w.get('value'), float) else str(w.get('value', '—')),
                     w.get('pass', '1'), str(w.get('element', '—'))]
                    for w in all_warnings]
            _data_table(doc,
                [['Design Warning', 'Limit', 'Actual Value', 'Pass', 'Element']],
                rows,
                col_widths=[7.5, 2.0, 2.0, 1.0, 2.0]
            )
        else:
            p = doc.add_paragraph()
            _run(p, 'no design warnings',
                 italic=True, color=BRAND_COLOR)

        _spacer(doc, 6)
        _heading(doc, 'System Saturation Indices (PHREEQC)')

        # Prefer the PHREEQC SI data attached during physics projection.
        # Fall back to the legacy heuristic if not available.
        conc_si = ro_results.get('concentrate_si')
        conc_ph = ro_results.get('concentrate_ph')
        feed_si = ro_results.get('feed_si', {})

        if conc_si:
            _spacer(doc, 4)

            SI_LIMITS = {
                'Calcite':   {'mod': 0.0,  'high': 0.5, 'crit': 1.0,  'formula': 'CaCO3',        'rec_mod': 'Acid dosing or antiscalant recommended.', 'rec_high': 'Antiscalant dosing required.', 'rec_crit': 'Critical: reduce recovery or use strong acid/antiscalant.'},
                'Aragonite': {'mod': 0.0,  'high': 0.5, 'crit': 1.0,  'formula': 'CaCO3 (orth.)', 'rec_mod': 'Carbonate scaling risk.', 'rec_high': 'Antiscalant dosing required.', 'rec_crit': 'Critical: carbonate scaling very likely.'},
                'Dolomite':  {'mod': 0.0,  'high': 1.0, 'crit': 2.0,  'formula': 'CaMg(CO3)2',   'rec_mod': 'Monitor dolomite saturation.', 'rec_high': 'Antiscalant recommended.', 'rec_crit': 'Severe dolomite scaling risk.'},
                'Gypsum':    {'mod': 0.0,  'high': 0.3, 'crit': 0.5,  'formula': 'CaSO4.2H2O',   'rec_mod': 'Antiscalant recommended.', 'rec_high': 'Antiscalant required; reduce recovery if possible.', 'rec_crit': 'Critical: Gypsum scale likely even with antiscalant.'},
                'Anhydrite': {'mod': 0.0,  'high': 0.3, 'crit': 0.5,  'formula': 'CaSO4',         'rec_mod': 'Antiscalant recommended.', 'rec_high': 'Antiscalant required.', 'rec_crit': 'Critical: reduce recovery.'},
                'Barite':    {'mod': -0.2, 'high': 0.0, 'crit': 0.3,  'formula': 'BaSO4',         'rec_mod': 'Specialized antiscalant needed.', 'rec_high': 'Antiscalant required - Barite is very insoluble.', 'rec_crit': 'Critical: Barite scale almost certain without inhibitor.'},
                'Celestite': {'mod': 0.0,  'high': 0.2, 'crit': 0.4,  'formula': 'SrSO4',         'rec_mod': 'Specialized antiscalant recommended.', 'rec_high': 'Antiscalant required.', 'rec_crit': 'Critical: Celestite scaling likely.'},
                'Fluorite':  {'mod': 0.0,  'high': 0.5, 'crit': 0.5,  'formula': 'CaF2',          'rec_mod': 'Monitor fluoride-calcium balance.', 'rec_high': 'Antiscalant required.', 'rec_crit': 'Critical: Fluorite scaling risk.'},
                'SiO2(a)':   {'mod': -0.1, 'high': 0.0, 'crit': 0.2,  'formula': 'SiO2 (am.)',    'rec_mod': 'pH adjust. or silica antiscalant.', 'rec_high': 'Silica antiscalant required.', 'rec_crit': 'Critical: Silica fouling very likely; reduce recovery.'},
            }

            RISK_COLORS = {
                'NONE':     (None, 'No action required.'),
                'LOW':      ('70CF9F', 'Monitor. No immediate action required.'),
                'MODERATE': ('F59E0B', None),
                'HIGH':     ('F97316', None),
                'CRITICAL': ('EF4444', None),
            }

            si_rows = []
            for mineral, si_val in conc_si.items():
                lim = SI_LIMITS.get(mineral)
                if not lim:
                    continue
                risk = 'NONE'
                rec  = 'No action required.'
                if si_val > lim['crit']:
                    risk = 'CRITICAL'; rec = lim['rec_crit']
                elif si_val > lim['high']:
                    risk = 'HIGH';     rec = lim['rec_high']
                elif si_val > lim['mod']:
                    risk = 'MODERATE'; rec = lim['rec_mod']
                elif si_val > lim['mod'] - 0.2:
                    risk = 'LOW';      rec = 'Monitor. No immediate action required.'
                
                f_si_val = feed_si.get(mineral)
                
                if si_val <= -99:
                    risk = 'NONE'
                    rec = 'Constituent ions not present.'
                    si_str = '—'
                else:
                    si_str = f'{si_val:+.3f}'
                    
                if f_si_val is not None and f_si_val > -99:
                    f_si_str = f'{f_si_val:+.3f}'
                else:
                    f_si_str = '—'
                    
                si_rows.append([mineral, lim['formula'], f_si_str, si_str, risk, rec])
            if conc_ph is not None:
                # Add Feed pH if available
                feed_ph_val = ro_results.get('feed', {}).get('ph') if 'feed' in ro_results else '—'
                si_rows.append(['pH', 'Equilibrium H+', f"{feed_ph_val:.2f}" if isinstance(feed_ph_val, float) else str(feed_ph_val), f'{conc_ph:.2f}', 'COMPUTED', 'Dynamically computed from feed HCO3.'])

            _data_table(doc,
                [['Mineral', 'Formula', 'Feed SI', 'Conc. SI', 'Risk', 'Recommendation']],
                si_rows,
                col_widths=[1.7, 2.0, 1.4, 1.4, 1.7, 5.8]
            )
        else:
            # Legacy fallback
            from system_engine import _compute_nf_concentrate_scaling
            sm = ro_results.get('summary', {})
            conc_ions = sm.get('conc_ions', {})
            feed_ph = ro_results.get('feed_ph', 7.0)
            temp_c = 25.0
            try:
                if 'feed' in ro_results and 'temperature' in ro_results['feed']:
                    temp_c = float(ro_results['feed']['temperature'])
            except Exception:
                pass
            sol_rows = []
            try:
                scaling = _compute_nf_concentrate_scaling(conc_ions, feed_ph, temp_c)
                if scaling.get('CaCO3_calcite', {}).get('antiscalant_required'):
                    sol_rows.append(['Calcite saturation exceeded (LSI > 0) - Risk of carbonate scaling.', 'Acid dosing or anti-scalant required.'])
                if scaling.get('CaSO4_gypsum', {}).get('antiscalant_required'):
                    sol_rows.append(['Gypsum saturation exceeded (SI > 0) - Risk of calcium sulfate scaling.', 'Verify anti-scalant efficacy.'])
                if scaling.get('BaSO4_barite', {}).get('antiscalant_required'):
                    sol_rows.append(['Barite saturation exceeded (SI > 0) - High risk of barium sulfate scaling.', 'Specialized anti-scalant required.'])
                if scaling.get('SrSO4_celestite', {}).get('antiscalant_required'):
                    sol_rows.append(['Celestite saturation exceeded (SI > 0) - Risk of strontium scaling.', 'Use specialized anti-scalant.'])
                if scaling.get('CaF2_fluorite', {}).get('antiscalant_required'):
                    sol_rows.append(['Fluorite saturation exceeded (SI > 0) - Risk of calcium fluoride scaling.', 'Use specialized anti-scalant.'])
                if scaling.get('SiO2', {}).get('antiscalant_required'):
                    sol_rows.append(['Amorphous Silica saturation > 80% - Risk of silica polymerization.', 'Consider pH adjustment or silica-specific anti-scalant.'])
            except Exception:
                pass
            if not sol_rows:
                sol_rows.append(['No critical solubility limits exceeded for major scalants.', 'General anti-scalant recommended as best practice.'])
            _data_table(doc, [['Warning', 'Recommended Solution']], sol_rows, col_widths=[7.5, 7.0])

        _spacer(doc, 6)

        _heading(doc, 'RO Chemical Adjustments (Pass 1 Feed vs Concentrate)')
        sm = ro_results.get('summary', {})
        _data_table(doc,
            [['Parameter', 'Pass 1 Feed', 'RO 1st Pass Conc.']],
            [
                ['pH',                          f"{ro_results.get('feed_ph', 7.0):.1f}",  '—'],
                ['TDS* (mg/L)',                 f"{sm.get('feed_tds', 0):.1f}", f"{sm.get('conc_tds', 0):.0f}"],
                ['Conductivity (µS/cm)',        f"{sm.get('feed_tds', 0)*2.02:.0f}", f"{sm.get('conc_tds', 0)*2.02:.0f}"],
                ['CaSO₄ (% saturation)',        '—', '—'],
                ['BaSO₄ (% saturation)',        '—', '—'],
                ['SrSO₄ (% saturation)',        '—', '—'],
                ['SiO₂ (% saturation)',         '—', '—'],
            ],
            col_widths=[7.0, 3.0, 4.5]
        )

    # ── Page 6: Economics ────────────────────────────────────────────────────

    def _page_economics(self, doc, economics):
        _heading(doc, 'Economic Analysis Summary', size=16)

        capex   = economics.get('capex', {})
        opex    = economics.get('opex', {})
        metrics = economics.get('metrics', {})

        # CAPEX table
        _heading(doc, 'Capital Expenditure (CAPEX)')
        capex_rows = [
            ['RO Membrane Elements',           f"₹ {capex.get('membranes_inr', 0):,.0f}"],
            ['Pressure Vessels',               f"₹ {capex.get('vessels_inr', 0):,.0f}"],
            ['High Pressure Pump',             f"₹ {capex.get('hp_pump_inr', 0):,.0f}"],
            ['Booster Pumps',                  f"₹ {capex.get('booster_pump_inr', 0):,.0f}"],
        ]
        if capex.get('uf_modules_inr', 0) > 0:
            capex_rows.append(['UF Membrane Modules',  f"₹ {capex.get('uf_modules_inr', 0):,.0f}  ({capex.get('uf_modules_count', 0)} modules)"])
            capex_rows.append(['UF Feed & Backwash Pumps', f"₹ {capex.get('uf_pumps_inr', 0):,.0f}"])
        capex_rows += [
            ['Equipment Subtotal',             f"₹ {capex.get('equip_subtotal_inr', 0):,.0f}"],
            ['Installation & Commissioning',   f"₹ {capex.get('ic_inr', 0):,.0f}"],
            ['Contingency',                    f"₹ {capex.get('contingency_inr', 0):,.0f}"],
            ['TOTAL CAPEX',                    f"₹ {capex.get('total_capex_inr', 0):,.0f}"],
        ]
        _data_table(doc,
            [['Item', 'Cost (INR)']],
            capex_rows,
            col_widths=[9.5, 5.0]
        )
        _spacer(doc, 6)

        # OPEX table
        _heading(doc, 'Annual Operating Expenditure (OPEX)')
        opex_rows = [
            ['Annual Operating Hours',       f"{opex.get('annual_hours', 0):,.0f} hrs"],
            ['Electricity / Energy Cost',    f"₹ {opex.get('energy_cost_pa_inr', 0):,.0f}"],
            ['RO Membrane Replacement',      f"₹ {opex.get('membrane_repl_pa_inr', 0):,.0f}"],
        ]
        if opex.get('uf_ceb_chemicals_pa_inr', 0) > 0:
            opex_rows.append(['UF CEB Chemicals (Citric Acid + NaOCl)', f"₹ {opex.get('uf_ceb_chemicals_pa_inr', 0):,.0f}"])
        opex_rows.append(['TOTAL OPEX', f"₹ {opex.get('total_opex_pa_inr', 0):,.0f}"])
        _data_table(doc,
            [['Item', 'Cost (INR/year)']],
            opex_rows,
            col_widths=[9.5, 5.0]
        )
        _spacer(doc, 6)

        # Financial metrics
        _heading(doc, 'Financial Indicators & Levelised Cost')
        _data_table(doc,
            [['Metric', 'Value']],
            [
                ['Capital Recovery Factor (CRF)',          f"{metrics.get('crf', 0):.4f}"],
                ['Annualised CAPEX',                       f"₹ {metrics.get('annualised_capex_inr', 0):,.0f} / year"],
                ['Total Annual Cost (CAPEX + OPEX)',        f"₹ {metrics.get('total_annual_cost_inr', 0):,.0f} / year"],
                ['Annual Water Production',                f"{metrics.get('annual_production_kl', 0):,.0f} KL / year"],
                ['Levelised Cost of Water (LCOW)',         f"₹ {metrics.get('cost_per_kl_inr', 0):.2f} / KL"],
            ],
            col_widths=[9.5, 5.0]
        )
        _spacer(doc, 8)

        # Economics charts side by side
        _heading(doc, '7 – CAPEX Breakdown & Annual OPEX')
        buf_pie = _chart_capex_pie(capex)
        buf_opex_bar = _chart_opex_bar(opex)
        _insert_two_charts(doc, buf_pie, buf_opex_bar, width_in=3.0)

    # ── Physics-Based Multi-Year Projection ────────────────────────────────────

    def _chart_npf_nsp_trend(self, snapshots, figsize=(5.5, 2.8)):
        """NPF and NSP trend over years."""
        years = [s['year'] for s in snapshots]
        npfs  = [s.get('npf', 1.0) for s in snapshots]
        nsps  = [s.get('nsp', 1.0) for s in snapshots]
        fig, ax1 = plt.subplots(figsize=figsize)
        ax2 = ax1.twinx()
        l1, = ax1.plot(years, npfs, marker='o', color=CHART_BRAND, linewidth=2, markersize=5, label='NPF')
        l2, = ax2.plot(years, nsps, marker='s', color=CHART_ORANGE, linewidth=2, markersize=5, linestyle='--', label='NSP')
        ax1.axhline(0.85, color='red', linestyle=':', linewidth=0.8, alpha=0.6)
        ax1.set_title('Normalised Performance Factors (ASTM D4516-19a)')
        ax1.set_xlabel('Year'); ax1.set_ylabel('NPF (Permeability)', color=CHART_BRAND)
        ax2.set_ylabel('NSP (Salt Passage)', color=CHART_ORANGE)
        ax1.set_ylim(bottom=min(0.5, min(npfs)*0.9), top=1.05)
        ax2.set_ylim(bottom=0.5, top=max(1.1, max(nsps)*1.1))
        ax1.set_xticks(years)
        ax1.legend(handles=[l1, l2], fontsize=7, loc='lower left')
        ax1.spines['top'].set_visible(False)
        fig.tight_layout()
        return _chart_bytes(fig)

    def _chart_fouling_mechanism_stacked(self, snapshots, figsize=(5.5, 2.8)):
        """Stacked area chart of fouling mechanism contributions."""
        years = [s['year'] for s in snapshots]
        rc  = [s.get('rc_avg', 0) * 1e-10 for s in snapshots]
        rb  = [s.get('rb_avg', 0) * 1e-10 for s in snapshots]
        rs  = [s.get('rs_avg', 0) * 1e-10 for s in snapshots]
        rn  = [s.get('rn_avg', 0) * 1e-10 for s in snapshots]
        rcp = [s.get('rcomp', 0) * 1e-10 for s in snapshots]
        fig, ax = plt.subplots(figsize=figsize)
        ax.stackplot(years, rc, rb, rs, rn, rcp,
                     labels=['Cake (Rc)', 'Biofilm (Rb)', 'Scale (Rs)', 'NOM (Rn)', 'Compaction'],
                     colors=[CHART_BLUE, '#E65100', '#9C27B0', '#009688', CHART_GRAY],
                     alpha=0.85)
        ax.set_title('Fouling Resistance Contributions Over Time')
        ax.set_xlabel('Year'); ax.set_ylabel('Resistance (×10¹⁰ m⁻¹)')
        ax.set_xticks(years)
        ax.legend(fontsize=6, loc='upper left')
        fig.tight_layout()
        return _chart_bytes(fig)

    def _chart_pressure_sec_trend(self, snapshots, figsize=(5.5, 2.6)):
        """Feed pressure and SEC over years."""
        years = [s['year'] for s in snapshots]
        pfeed = [s.get('feed_pressure_bar', 0) for s in snapshots]
        sec   = [s.get('sec_kwh_m3', 0) for s in snapshots]
        fig, ax1 = plt.subplots(figsize=figsize)
        ax2 = ax1.twinx()
        l1, = ax1.plot(years, pfeed, marker='o', color=CHART_BRAND, linewidth=2, markersize=5, label='Feed Pressure (bar)')
        l2, = ax2.plot(years, sec, marker='s', color=CHART_ORANGE, linewidth=2, markersize=5, linestyle='--', label='SEC (kWh/m³)')
        ax1.set_title('Feed Pressure & Specific Energy vs Year')
        ax1.set_xlabel('Year'); ax1.set_ylabel('Feed Pressure (bar)', color=CHART_BRAND)
        ax2.set_ylabel('SEC (kWh/m³)', color=CHART_ORANGE)
        ax1.set_xticks(years)
        ax1.legend(handles=[l1, l2], fontsize=7, loc='upper left')
        ax1.spines['top'].set_visible(False)
        fig.tight_layout()
        return _chart_bytes(fig)

    def _page_physics_projection(self, doc, physics: dict, proj_name: str,
                                  selected_year: int = 0):
        """Generate the Physics-Based Multi-Year Performance Projection page."""
        _heading(doc, 'Year-Wise Performance', size=16)
        _spacer(doc, 6)

        snapshots = physics.get('annual_snapshots', [])
        if not snapshots:
            _run(doc.add_paragraph(), 'No physics projection data available.', size=9)
            return

        # ── Selected-year highlighted summary ────────────────────────────────
        sel = next((s for s in snapshots if s['year'] == selected_year), snapshots[-1])
        base = snapshots[0]

        _heading(doc, f'Year {selected_year} Performance Summary (vs Year 0 Baseline)', size=11)
        _data_table(doc,
            [['Parameter', 'Year 0 (Baseline)', f'Year {selected_year}', 'Change']],
            [
                ['Feed Pressure',
                 f"{base.get('feed_pressure_bar', 0):.2f} bar",
                 f"{sel.get('feed_pressure_bar', 0):.2f} bar",
                 f"{sel.get('feed_pressure_bar', 0) - base.get('feed_pressure_bar', 0):+.2f} bar"],
                ['Permeate Flow',
                 f"{base.get('perm_flow', 0):.2f} m\u00b3/h",
                 f"{sel.get('perm_flow', 0):.2f} m\u00b3/h",
                 f"{sel.get('perm_flow', 0) - base.get('perm_flow', 0):+.2f} m\u00b3/h"],
                ['Recovery',
                 f"{base.get('recovery', 0)*100:.1f}%",
                 f"{sel.get('recovery', 0)*100:.1f}%",
                 f"{(sel.get('recovery', 0) - base.get('recovery', 0))*100:+.1f}%"],
                ['Permeate TDS',
                 f"{base.get('perm_tds', 0):.1f} mg/L",
                 f"{sel.get('perm_tds', 0):.1f} mg/L",
                 f"{sel.get('perm_tds', 0) - base.get('perm_tds', 0):+.1f} mg/L"],
                ['SEC',
                 f"{base.get('sec_kwh_m3', 0):.3f} kWh/m\u00b3",
                 f"{sel.get('sec_kwh_m3', 0):.3f} kWh/m\u00b3",
                 f"{sel.get('sec_kwh_m3', 0) - base.get('sec_kwh_m3', 0):+.3f} kWh/m\u00b3"],
                ['NPF (Norm. Permeate Flow)',
                 f"{base.get('npf', 1.0):.3f}",
                 f"{sel.get('npf', 1.0):.3f}",
                 f"{sel.get('npf', 1.0) - 1.0:+.3f}"],
                ['NSP (Norm. Salt Passage)',
                 f"{base.get('nsp', 1.0):.3f}",
                 f"{sel.get('nsp', 1.0):.3f}",
                 f"{sel.get('nsp', 1.0) - 1.0:+.3f}"],
                ['FRI (Fouling Resistance Index)',
                 '0.000',
                 f"{sel.get('fri', 0):.4f}",
                 f"{sel.get('fri', 0):+.4f}"],
                ['B Relative (Salt Perm. Degradation)',
                 '1.000',
                 f"{sel.get('b_irr', 1.0):.4f}",
                 f"{sel.get('b_irr', 1.0) - 1.0:+.4f}"],
            ],
            col_widths=[5.5, 3.2, 3.2, 2.6]
        )
        _spacer(doc, 8)

        # ── Year-by-year performance table ────────────────────────────────────
        _heading(doc, 'Year-by-Year Performance Table', size=11)
        table_rows = []
        for s in snapshots:
            cip_flag = 'YES' if s.get('cip_triggered') else ''
            repl_flag = 'YES' if s.get('replacement_triggered') else ''
            table_rows.append([
                str(s['year']),
                f"{s.get('perm_flow', 0):.2f}",
                f"{s.get('recovery', 0)*100:.1f}",
                f"{s.get('feed_pressure_bar', 0):.2f}",
                f"{s.get('perm_tds', 0):.1f}",
                f"{s.get('sec_kwh_m3', 0):.3f}",
                f"{s.get('npf', 1.0):.3f}",
                f"{s.get('nsp', 1.0):.3f}",
                f"{s.get('fri', 0):.4f}",
                cip_flag,
                repl_flag,
            ])
        _data_table(doc,
            [['Year', 'Qp\n(m\u00b3/h)', 'Rec.\n(%)', 'P_feed\n(bar)', 'TDS\n(mg/L)',
              'SEC\n(kWh/m\u00b3)', 'NPF', 'NSP', 'FRI', 'CIP', 'Repl.']],
            table_rows,
            col_widths=[1.0, 1.5, 1.3, 1.5, 1.5, 1.8, 1.3, 1.3, 1.5, 1.0, 1.0]
        )
        _spacer(doc, 8)

        # ── Fouling mechanism breakdown at selected year ───────────────────────
        _heading(doc, f'Fouling Resistance Breakdown at Year {selected_year}', size=11)
        rf_sum = sel.get('rc_avg', 0) + sel.get('rb_avg', 0) + sel.get('rs_avg', 0) + sel.get('rn_avg', 0)
        _data_table(doc,
            [['Fouling Mechanism', 'Resistance (m\u207b\u00b9)', 'Relative Contribution']],
            [
                ['Colloidal Cake (Rc)', f"{sel.get('rc_avg', 0):.3e}",
                 f"{sel.get('rc_avg', 0) / max(rf_sum, 1e-10) * 100:.1f}%"],
                ['Biofilm (Rb)', f"{sel.get('rb_avg', 0):.3e}",
                 f"{sel.get('rb_avg', 0) / max(rf_sum, 1e-10) * 100:.1f}%"],
                ['Mineral Scaling (Rs)', f"{sel.get('rs_avg', 0):.3e}",
                 f"{sel.get('rs_avg', 0) / max(rf_sum, 1e-10) * 100:.1f}%"],
                ['NOM Adsorption (Rn)', f"{sel.get('rn_avg', 0):.3e}",
                 f"{sel.get('rn_avg', 0) / max(rf_sum, 1e-10) * 100:.1f}%"],
                ['Compaction (Structural)', f"{sel.get('rcomp', 0):.3e}", 'N/A'],
            ],
            col_widths=[5.5, 4.5, 4.5]
        )
        _spacer(doc, 6)

        # ── Scaling indices at wall (Year N) ──────────────────────────────────
        _heading(doc, f'Wall-Level Saturation Indices at Year {selected_year}', size=11)
        p = doc.add_paragraph()
        _run(p, 'Note: Wall-SI includes concentration polarisation enhancement — always higher than bulk concentrate SI.',
             size=7.5, italic=True, color=RGBColor(100,116,139))
        _data_table(doc,
            [['Mineral', 'Bulk SI (Feed)', 'Wall SI (Year 0)', f'Wall SI (Year {selected_year})', 'Risk']],
            [
                ['Calcite (CaCO\u2083)',
                 f"{base.get('si_calcite_wall', 0):.2f}",
                 f"{base.get('si_calcite_wall', 0):.2f}",
                 f"{sel.get('si_calcite_wall', 0):.2f}",
                 'HIGH' if sel.get('si_calcite_wall', 0) > 0 else 'LOW'],
                ['Gypsum (CaSO\u2084)',
                 f"{base.get('si_gypsum_wall', 0):.2f}",
                 f"{base.get('si_gypsum_wall', 0):.2f}",
                 f"{sel.get('si_gypsum_wall', 0):.2f}",
                 'HIGH' if sel.get('si_gypsum_wall', 0) > 0 else 'LOW'],
                ['Barite (BaSO\u2084)',
                 f"{base.get('si_barite_wall', 0):.2f}",
                 f"{base.get('si_barite_wall', 0):.2f}",
                 f"{sel.get('si_barite_wall', 0):.2f}",
                 'HIGH' if sel.get('si_barite_wall', 0) > 0 else 'LOW'],
                ['Silica (SiO\u2082)',
                 f"{base.get('si_silica_wall', 0):.2f}",
                 f"{base.get('si_silica_wall', 0):.2f}",
                 f"{sel.get('si_silica_wall', 0):.2f}",
                 'HIGH' if sel.get('si_silica_wall', 0) > 0 else 'LOW'],
            ],
            col_widths=[3.5, 2.8, 3.0, 3.0, 2.2]
        )
        _spacer(doc, 8)


        # ── Charts ────────────────────────────────────────────────────────────
        _heading(doc, 'Performance Trend Charts', size=11)
        buf_npf = self._chart_npf_nsp_trend(snapshots)
        buf_prs = self._chart_pressure_sec_trend(snapshots)
        _insert_two_charts(doc, buf_npf, buf_prs, width_in=3.0)
        _spacer(doc, 4)

        buf_mech = self._chart_fouling_mechanism_stacked(snapshots)
        if buf_mech:
            _insert_chart(doc, buf_mech, width_in=5.5)

        _spacer(doc, 8)

    def _page_aging_summary(self, doc, aging: dict, proj_name: str):
        _heading(doc, 'Membrane Aging Summary', size=16)
        _spacer(doc, 6)

        profile = aging.get('aging_profile', [])
        final_state = profile[-1] if profile else {}

        _data_table(doc,
            [['Parameter', 'Value']],
            [
                ['End of Life (Month)', f"Month {aging.get('end_of_life_month', 'N/A')}"],
                ['Dominant Degradation Mechanism', str(aging.get('dominant_mechanism', 'N/A')).replace('_', ' ').title()],
                ['Final Normalized Permeate Flow (NPF)', f"{final_state.get('npf', 0):.3f}"],
                ['Final Feed Pressure', f"{final_state.get('p_feed_bar', 0):.1f} bar"],
                ['Status', str(aging.get('status', 'N/A')).title()],
            ],
            col_widths=[7.5, 7.0]
        )
        _spacer(doc, 8)

    # ── Footnote ─────────────────────────────────────────────────────────────

    def _footnote(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        _run(p, 'Footnotes: ', bold=True, size=8)
        _run(p, '*TDS and Conductivity include dissolved ionic salts only (excludes SiO₂, B, CO₂). ',
             italic=True, size=8)

    # ══════════════════════════════════════════════════════════════════════════
    # Public entry point
    # ══════════════════════════════════════════════════════════════════════════

    def generate_calculation_report(self, system_results: dict,
                                    output_path: str = 'Calculation_Report.docx') -> str:
        doc   = self._setup_doc()
        sr    = system_results
        train = sr.get('technology_train', '1P-RO')
        proj  = sr.get('project_name', 'PACE Report')
        fw    = sr.get('feed_water_used', {})

        ro_results   = sr.get('ro_results')
        pass1_res    = sr.get('pass1_results')
        pass2_res    = sr.get('pass2_results')
        ro_main      = pass1_res if pass1_res else ro_results
        mem1         = sr.get('ro_membrane', '—')
        mem2         = sr.get('pass2_membrane', mem1)

        # ── Page 1 ────────────────────────────────────────────────────────────
        if sr.get('uf_results') and 'UF' in train:
            self._page_uf_overview(doc, sr)

        self._page_system_overview(doc, sr)

        # ── Page 2: Stage + Element tables ───────────────────────────────────
        if ro_main:
            self._page_flow_tables(doc, ro_main, mem1, pass_label='Pass 1', sr=sr)

        if pass2_res:
            self._page_flow_tables(doc, pass2_res, mem2, pass_label='Pass 2', sr=sr)

        # ── Page 3: Ion analysis ──────────────────────────────────────────────
        if ro_main:
            self._page_ion_analysis(doc, ro_main, fw)

        # ── Page 5: Warnings ──────────────────────────────────────────────────
        # Inject PHREEQC SI data into ro_main so _page_warnings can display it.
        # These keys are attached by the physics endpoint; missing on plain calculate-system calls.
        if ro_main is not None:
            ro_main['concentrate_si'] = sr.get('concentrate_si')
            ro_main['concentrate_ph'] = sr.get('concentrate_ph')
            ro_main['feed_si'] = sr.get('feed_si')
        self._page_warnings(doc, ro_main or {}, pass2_res)

        # ── Page 6: Economics ────────────────────────────────────────────────
        economics = sr.get('economics')
        if economics:
            self._page_economics(doc, economics)

        # ── Page 7: Membrane Aging ───────────────────────────────────────────
        aging_res = sr.get('aging_results')
        if aging_res:
            self._page_aging_summary(doc, aging_res, proj)

        # ── Page 8: Physics-Based Multi-Year Performance Projection ──────────
        physics_res = sr.get('physics_results')
        if physics_res:
            selected_yr = sr.get('physics_selected_year', 0)
            self._page_physics_projection(doc, physics_res, proj, selected_year=selected_yr)

        # Footer
        _add_footer(doc, proj, self.today)

        doc.save(output_path)
        plt.close('all')
        return output_path
