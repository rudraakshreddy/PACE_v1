from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Title
title = doc.add_heading('Scaling Analysis Module: Technical & Code Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('This document outlines the technical implementation, data flow, and code architecture of the PHREEQC-integrated Scaling Analysis Module.')

# 1. Architecture & Stack
doc.add_heading('1. Tech Stack & Architecture', level=1)
doc.add_paragraph(
    'The module relies on a decoupled client-server architecture:\n'
    '• Frontend: Vanilla JavaScript (script.js) with HTML5/CSS3 for UI.\n'
    '• Backend API: FastAPI (Python 3) handling asynchronous HTTP POST requests.\n'
    '• Thermodynamic Engine: phreeqpython wrapper over the USGS PHREEQC C++ engine.'
)

# 2. Frontend: Data Extraction & API Call
doc.add_heading('2. Frontend: Data Extraction (script.js)', level=1)
doc.add_paragraph(
    'In script.js, the function runPhreeqcCalculation() is triggered. It extracts the raw values from the HTML input fields.'
)
code_frontend = doc.add_paragraph(
    "const payload = {\n"
    "    temperature: parseFloat(document.getElementById('temp').value) || 25.0,\n"
    "    ph: parseFloat(document.getElementById('ph').value) || 7.0,\n"
    "    calcium: parseFloat(document.getElementById('ca').value) || 0,\n"
    "    ...\n"
    "    bicarbonate: parseFloat(document.getElementById('hco3').value) || 0\n"
    "};"
)
code_frontend.style = 'Intense Quote'
doc.add_paragraph('This payload is then sent to the backend via a fetch() POST request to http://localhost:8000/api/calculate-scaling.')

# 3. Backend: API Endpoint & Pydantic Schema
doc.add_heading('3. Backend: FastAPI Server (server.py)', level=1)
doc.add_paragraph(
    'The backend validates the incoming JSON using a Pydantic model (FeedWaterData). This ensures strict typing before passing '
    'data to the C++ engine.'
)
code_backend = doc.add_paragraph(
    "class FeedWaterData(BaseModel):\n"
    "    temperature: float\n"
    "    ph: float\n"
    "    calcium: float\n"
    "    bicarbonate: float\n"
    "    ...\n\n"
    "@app.post('/api/calculate-scaling')\n"
    "async def calculate_scaling(data: FeedWaterData):"
)
code_backend.style = 'Intense Quote'

# 4. Thermodynamic Mapping (The Core Logic)
doc.add_heading('4. PHREEQC Integration & Thermodynamic Mapping', level=1)
doc.add_paragraph(
    'Inside the endpoint, a PhreeqPython instance is created. The data is mapped to a dictionary representing an aqueous solution. '
    'Crucially, elements must be mapped to specific oxidation states or compound forms to ensure accurate molar mass calculation by PHREEQC.'
)
code_phreeqc = doc.add_paragraph(
    "solution_data = {\n"
    "    'units': 'mg/L',\n"
    "    'temp': data.temperature,\n"
    "    'pH': data.ph,\n"
    "    'Ca': data.calcium,\n"
    "    'S(6)': f\"{data.sulfate} as SO4\",\n"
    "    'C(4)': f\"{data.bicarbonate + data.carbonate} as CaCO3\"\n"
    "}\n"
    "sol = pp.add_solution(solution_data)"
)
code_phreeqc.style = 'Intense Quote'
doc.add_paragraph(
    'Important Logic Note: The Bicarbonate input is mapped to Total Inorganic Carbon (C(4)) using the "as CaCO3" suffix. '
    'This explicitly tells PHREEQC to treat the input value as Total Alkalinity as CaCO3 (the industry standard). '
    'This allows PHREEQC to accurately simulate the LSI after acid dosing by locking the Total Carbon mass, directly matching '
    'results from proprietary RO software like Genesys.'
)

# 5. Extraction and Response
doc.add_heading('5. Data Extraction and JSON Response', level=1)
doc.add_paragraph(
    'Once the solution is balanced, the sol.si() method extracts the exact Logarithmic Saturation Index for specific minerals.'
)
code_response = doc.add_paragraph(
    "return {\n"
    "    'gypsum_si': sol.si('Gypsum'),\n"
    "    'calcite_si': sol.si('Calcite'),\n"
    "    'silica_si': sol.si('SiO2(a)'),\n"
    "    'fluorite_si': sol.si('Fluorite'),\n"
    "    'lsi': sol.si('Calcite')  # Standard thermodynamic LSI proxy\n"
    "}"
)
code_response.style = 'Intense Quote'

# 6. Frontend: UI Rendering
doc.add_heading('6. Frontend UI Updates (script.js)', level=1)
doc.add_paragraph(
    'The frontend receives the JSON response and updates the glassmorphism UI cards. To show the % Saturation, it uses '
    'the standard mathematical transformation for Saturation Ratios:'
)
code_ui = doc.add_paragraph(
    "const pct = Math.pow(10, result.calcite_si) * 100;\n"
    "element.textContent = ${result.calcite_si.toFixed(3)} (%);"
)
code_ui.style = 'Intense Quote'

doc.save('PHREEQC_Code_Documentation.docx')
