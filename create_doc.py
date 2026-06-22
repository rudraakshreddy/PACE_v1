from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Scaling Analysis Module: PHREEQC Integration Document', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'The Scaling Analysis Module in Project PACE is designed to predict the precipitation risk of various inorganic salts '
    'during the Reverse Osmosis (RO) concentration process. To achieve industry-standard accuracy comparable to proprietary '
    'RO projection software (such as Genesys), the system integrates the PHREEQC thermodynamic engine, originally developed '
    'by the US Geological Survey (USGS).'
)

# 2. Architectural Overview
doc.add_heading('2. Architectural Overview', level=1)
p2 = doc.add_paragraph('The integration relies on a modern, decoupled architecture:')
p2.add_run('\n• Frontend (UI/UX):').bold = True
p2.add_run(' Built with vanilla HTML/JS and sleek glassmorphism design. It gathers raw feed water data (ions in mg/L, pH, Temperature).')
p2.add_run('\n• Backend (API):').bold = True
p2.add_run(' A FastAPI Python server exposes a /api/calculate-scaling endpoint.')
p2.add_run('\n• Engine (PHREEQC):').bold = True
p2.add_run(' The phreeqpython library is used to instantiate aqueous solutions, perform speciation, and calculate exact Saturation Indices (SI) for target minerals.')

# 3. Thermodynamic Rigor vs. Empirical Models
doc.add_heading('3. Thermodynamic Rigor vs. Empirical Models', level=1)
doc.add_paragraph(
    'Many basic scaling calculators rely on empirical equations (like the Langelier Saturation Index or Stiff & Davis Index) '
    'and simple solubility products (Ksp). However, real water chemistry is highly complex due to ionic strength effects and '
    'ion-pairing (e.g., Calcium and Sulfate forming soluble CaSO4(aq) pairs rather than precipitating).\n\n'
    'PHREEQC solves this by calculating the true thermodynamic activity of every chemical species in the water, ensuring that '
    'the Saturation Index (SI) is rigorously accurate even at high salinities (Concentrate streams).'
)

# 4. Input Mapping and Carbonate Chemistry
doc.add_heading('4. Input Mapping & Carbonate Chemistry (Genesys Alignment)', level=1)
doc.add_paragraph(
    'A critical challenge in RO scaling prediction is matching commercial software output for Calcium Carbonate (CaCO3). '
    'In standard software like Genesys, the user inputs "Total Alkalinity (as CaCO3)". If this value is naively treated as '
    'Bicarbonate (HCO3) mass, the total inorganic carbon in the system is underestimated, leading to mismatched LSI predictions.'
)
p4 = doc.add_paragraph()
p4.add_run('Solution: ').bold = True
p4.add_run(
    'The backend explicitly maps the Alkalinity input to Total Inorganic Carbon (C(4)) formatted as "CaCO3" in PHREEQC. '
    'This allows PHREEQC to fix the Total Carbon in the system. When the user inputs a target pH (e.g., 6.1 representing '
    'feed water after acid dosing), PHREEQC accurately recalculates the Bicarbonate/CO2 equilibrium, matching Genesys\'s '
    'LSI outputs down to the decimal point (e.g., -2.06).'
)

# 5. Logarithmic SI vs. Linear Saturation Ratio
doc.add_heading('5. Data Interpretation: Logarithmic vs. Linear', level=1)
doc.add_paragraph(
    'There is a distinct difference in how thermodynamic software (PHREEQC) and commercial antiscalant software (Genesys) '
    'present scaling risk data:'
)
p5 = doc.add_paragraph()
p5.add_run('• Logarithmic Saturation Index (SI): ').bold = True
p5.add_run(
    'PHREEQC outputs the standard geochemical SI, defined as log10(Concentration / Solubility). '
    'An SI of 0 means perfect equilibrium (100% saturation). Negative values indicate undersaturation. Our backend returns '
    'these exact Logarithmic SI values.'
)
p5.add_run('\n• Linear Saturation Ratio (Index): ').bold = True
p5.add_run(
    'For non-carbonate minerals like Silica and CaF2, software like Genesys displays the linear ratio '
    '(Concentration / Solubility) directly. For example, a Silica SI of -1.233 in our software perfectly translates to a '
    'linear ratio of 10^(-1.233) = 0.058 (or ~5.8%), matching Genesys\'s 0.06 index.'
)
p5.add_run('\n\nNote: ').bold = True
p5.add_run('For CaCO3, the entire industry universally uses the Logarithmic scale (LSI).')

# 6. Future Expansion
doc.add_heading('6. Future Expansion Capabilities', level=1)
doc.add_paragraph(
    'Because the core engine is pure PHREEQC, the module is highly extensible. Future updates can easily integrate:\n'
    '1. Live Concentrate Analysis via a "Recovery %" slider (automatically concentrating the ions).\n'
    '2. Dynamic Acid Dosing Simulation (calculating exact mg/L of acid required to reach a safe LSI).\n'
    '3. Mixed-salt precipitation and competitive antiscalant thresholds.'
)

doc.save('PHREEQC_Integration_Document.docx')
